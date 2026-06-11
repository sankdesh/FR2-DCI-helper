# -*- coding: utf-8 -*-
"""Build 3GPP_SPEC_TRACEABILITY.docx from 3GPP_SPEC_TRACEABILITY.md (python-docx).

Run from repo root or Try 2:
  python Try 2/build_spec_traceability_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_paragraph_with_bold(doc: Document, text: str, style: str | None = None) -> None:
    """Split on **bold** and add runs."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [label](url) -> label
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    p = doc.add_paragraph(style=style)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            p.add_run(part)


def _is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().split("|")[1:-1]]
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c) for c in cells)


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().split("|")[1:-1]]


def _plain_for_cell(val: str) -> str:
    val = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", val)
    val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
    return val


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.style = "Normal"
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip().startswith("|") and "|" in line[1:]:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i]
                if _is_table_separator(row_line):
                    i += 1
                    continue
                rows.append(_parse_table_row(row_line))
                i += 1
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            # pad rows
            for r in rows:
                while len(r) < ncols:
                    r.append("")
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"
            for ri, row_cells in enumerate(rows):
                for ci, val in enumerate(row_cells):
                    tbl.rows[ri].cells[ci].text = _plain_for_cell(val)
            if rows:
                for ci in range(ncols):
                    _set_cell_shading(tbl.rows[0].cells[ci], "E7E6E6")
            if rows:
                for ci in range(ncols):
                    _set_cell_shading(tbl.rows[0].cells[ci], "E7E6E6")
            doc.add_paragraph()
            continue
        elif not line.strip():
            doc.add_paragraph()
        else:
            _add_paragraph_with_bold(doc, line)

        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> int:
    base = Path(__file__).resolve().parent
    md = base / "3GPP_SPEC_TRACEABILITY.md"
    out = base / "3GPP_SPEC_TRACEABILITY.docx"
    if not md.is_file():
        print(f"ERROR: missing {md}", file=sys.stderr)
        return 1
    convert_md_to_docx(md, out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
